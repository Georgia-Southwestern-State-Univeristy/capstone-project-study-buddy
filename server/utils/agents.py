""" This module contains the agent functions that interact with the external APIs. """
"""STEP 1: Import necessary modules"""
import os
import random
import requests
from langchain_google_community import GoogleSearchAPIWrapper
from langchain_community.utilities import BingSearchAPIWrapper
from langchain_community.tools import YouTubeSearchTool
from azure.core.credentials import AzureKeyCredential
from azure.ai.textanalytics import TextAnalyticsClient
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document as DocxDocument
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from uuid import uuid4
import os
from PIL import Image, ImageDraw, ImageFont
import uuid
import time

# Initialize Azure Text Analytics Client
text_analytics_key = os.getenv("AZURE_TEXT_ANALYTICS_KEY")
text_analytics_endpoint = os.getenv("AZURE_TEXT_ANALYTICS_ENDPOINT")
text_analytics_client = TextAnalyticsClient(endpoint=text_analytics_endpoint, credential=AzureKeyCredential(text_analytics_key))


"""Step 2: Define the agent functions"""

def get_google_search_results(query: str):
    """
    Uses Google Custom Search to fetch search results for a given query.

    Args:
        query (str): The search query.

    Returns:
        list: A list of search results with titles and links.
    """

    try:
        google_search_wrapper = GoogleSearchAPIWrapper(k=3)
        search_results = google_search_wrapper.run(query)
        print("Search results obtained:", search_results)

        # Ensure the results are JSON-serializable
        
        return search_results
    
    except Exception as e:
        print(f"Failed to fetch Google search results: {e}", exc_info=True)
        return None
    



    
def get_youtube_search_results(query: str):
    """
    Uses YouTube Search to fetch search results for a given query.

    Args:
        query (str): The search query.

    Returns:
        list: A list of search results with titles, descriptions, and video links.
    """
    try:
        youtube_search_tool = YouTubeSearchTool()
        search_results = youtube_search_tool.run(query)
        print("Search results obtained:", search_results)

        # Ensure the results are JSON-serializable
        return search_results

    except Exception as e:
        print(f"Failed to fetch YouTube search results: {e}", exc_info=True)
        return None





def get_bing_search_results(query: str):
    """
    Uses Bing Search to fetch search results for a given query.

    Args:
        query (str): The search query.

    Returns:
        list: A list of search results with titles and links.
    """
    try:
        bing_search_wrapper = BingSearchAPIWrapper()
        search_results = bing_search_wrapper.run(query)
        print("Search results obtained:", search_results)

        # Ensure the results are JSON-serializable
        return search_results

    except Exception as e:
        print(f"Failed to fetch Bing search results: {e}", exc_info=True)
        return None
    




def generate_suggestions(mood: str, user_input: str):
    """
    Generates personalized activities or coping mechanisms based on the user's mood and sentiment using a language model.
    
    Args:
        mood (str): The user's current mood.
        user_input (str): The user's input text to analyze sentiment.
    
    Returns:
        list: A list of suggested activities or coping mechanisms.
    """
    # Analyze sentiment
    sentiment_response = text_analytics_client.analyze_sentiment(documents=[{"id": "1", "text": user_input}])
    sentiment = sentiment_response[0].sentiment

    # Generate suggestions based on mood and sentiment
    prompt = f"Suggest some personalized activities or coping mechanisms for someone who is feeling {mood} and has a sentiment of {sentiment}."
    response = text_analytics_client.analyze_sentiment(documents=[{"id": "2", "text": prompt}])
    suggestions = response[0].sentiment.split('\n')
    
    return suggestions


def get_public_domain_textbooks(query: str):
    """
    Searches for textbooks in public domain libraries based on the user's query.

    Args:
        query (str): The search query.

    Returns:
        str: A formatted string containing the search results with PDF links if available.
    """
    try:
        # Use Open Library Search API
        search_response = requests.get(
            "https://openlibrary.org/search.json",
            params={"title": query, "has_fulltext": "true"}
        )
        search_data = search_response.json()
        books = search_data.get("docs", [])[:3]  # Get top 3 results

        if not books:
            return "No textbooks found for your query."

        results = "Here are some textbooks you might find useful:\n"
        for book in books:
            title = book.get("title", "Unknown Title")
            author = ', '.join(book.get("author_name", ["Unknown Author"]))
            work_key = book.get('key')
            edition_key = book.get('edition_key', [None])[0]

            # Initialize PDF link
            pdf_link = None

            if edition_key:
                # Fetch edition data to check for available formats
                edition_response = requests.get(f"https://openlibrary.org/books/{edition_key}.json")
                edition_data = edition_response.json()
                formats = edition_data.get('formats', {})

                # Check if a PDF is available in formats
                if 'pdf' in formats:
                    pdf_link = formats['pdf'].get('url')

                # Alternatively, check for Internet Archive links
                elif 'ocaid' in edition_data:
                    ocaid = edition_data['ocaid']
                    pdf_link = f"https://archive.org/download/{ocaid}/{ocaid}.pdf"

            # Fallback to the work link if no PDF is available
            if pdf_link:
                link = pdf_link
                link_text = "Download PDF"
            else:
                link = f"https://openlibrary.org{work_key}"
                link_text = "Read online"

            results += f"- {title} by {author}\n  {link_text}: {link}\n"

        return results

    except Exception as e:
        print(f"Failed to fetch textbooks: {e}")
        return "Sorry, I couldn't fetch textbooks at the moment."
    

def get_gutendex_domain_textbooks(query: str):
    """
    Searches for textbooks in Project Gutenberg based on the user's query.

    Args:
        query (str): The search query.

    Returns:
        str: A formatted string containing the search results with download links.
    """
    try:
        # Use Project Gutenberg's catalog via a third-party API
        search_response = requests.get(
            "http://gutendex.com/books",
            params={"search": query}
        )
        search_data = search_response.json()
        books = search_data.get("results", [])[:3]  # Get top 3 results

        if not books:
            return "No textbooks found for your query."

        results = "Here are some textbooks you might find useful:\n"
        for book in books:
            title = book.get("title", "Unknown Title")
            authors = [author.get("name", "Unknown Author") for author in book.get("authors", [])]
            author = ', '.join(authors) if authors else "Unknown Author"
            formats = book.get("formats", {})
            pdf_link = formats.get("application/pdf")
            epub_link = formats.get("application/epub+zip")
            txt_link = formats.get("text/plain; charset=utf-8")

            # Provide available formats
            results += f"- {title} by {author}\n"
            if pdf_link:
                results += f"  Download PDF: {pdf_link}\n"
            if epub_link:
                results += f"  Download EPUB: {epub_link}\n"
            if txt_link:
                results += f"  Download TXT: {txt_link}\n"

        return results

    except Exception as e:
        print(f"Failed to fetch textbooks: {e}")
        return "Sorry, I couldn't fetch textbooks at the moment."



def generate_document(content: str, format: str = 'pdf'):
    """
    Generates a document with the given content and format.

    Args:
        content (str): The content to include in the document.
        format (str): The format of the document ('pdf' or 'docx').

    Returns:
        str: A message containing the download URL for the generated document.
    """
    # Generate a unique filename
    filename = f"{uuid4()}.{format}"
    file_path = os.path.join("generated_documents", filename)

    # Ensure the directory exists
    os.makedirs("generated_documents", exist_ok=True)

    if format == 'pdf':
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        flowables = []

        # Process content
        lines = content.split('\n')
        for line in lines:
            if line.startswith('## '):
                heading = Paragraph(line[3:], styles['Heading2'])
                flowables.append(heading)
            elif line.startswith('# '):
                heading = Paragraph(line[2:], styles['Heading1'])
                flowables.append(heading)
            elif line.startswith('- '):
                bullet = Paragraph(f"&bull; {line[2:]}", styles['BodyText'])
                flowables.append(bullet)
            else:
                paragraph = Paragraph(line, styles['BodyText'])
                flowables.append(paragraph)
            flowables.append(Spacer(1, 0.2 * inch))

        doc.build(flowables)
    elif format == 'docx':

        doc = DocxDocument()
        lines = content.split('\n')
        for line in lines:
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        doc.save(file_path)
    else:
        return "Unsupported format. Please choose 'pdf' or 'docx'."
    
    backend_base_url = os.getenv('BACKEND_BASE_URL', 'http://localhost:8000')  # Ensure this environment variable is set
    # Return the URL to download the file
    download_url = f"{backend_base_url}/ai_mentor/download_document/{filename}"
    print(f"Document generated: {download_url}")
    return download_url


def get_job_listings(skills: str, location: str = None):
    """
    Fetches job listings that match the user's skills.

    Args:
        skills (str): A comma-separated string of user skills.
        location (str, optional): The location to search for jobs.

    Returns:
        str: A formatted string containing the job listings.
    """
    # Replace with your Adzuna API credentials
    ADZUNA_APP_ID = os.getenv('ADZUNA_APP_ID')
    ADZUNA_APP_KEY = os.getenv('ADZUNA_APP_KEY')
    if not ADZUNA_APP_ID or not ADZUNA_APP_KEY:
        return "Job search API credentials are not set."

    # Prepare the API endpoint
    country = 'us'  # Change to your target country code
    base_url = f'https://api.adzuna.com/v1/api/jobs/{country}/search/1'
    params = {
        'app_id': ADZUNA_APP_ID,
        'app_key': ADZUNA_APP_KEY,
        'what': skills,
        'content-type': 'application/json',
    }
    if location:
        params['where'] = location

    try:
        response = requests.get(base_url, params=params)
        data = response.json()
        results = data.get('results', [])
        if not results:
            return "No job listings found matching your skills."

        # Format the job listings
        job_listings = "Here are some job listings matching your skills:\n"
        for job in results[:5]:  # Limit to top 5 results
            title = job.get('title', 'No title')
            company = job.get('company', {}).get('display_name', 'Unknown company')
            job_location = job.get('location', {}).get('display_name', 'Unknown location')
            url = job.get('redirect_url', '')
            job_listings += f"- **{title}** at **{company}** in **{job_location}**\n"
            job_listings += f"  [View Job Posting]({url})\n\n"
        return job_listings

    except Exception as e:
        print(f"Failed to fetch job listings: {e}")
        return "Sorry, I couldn't fetch job listings at the moment."
    

def fetch_meme(topic: str) -> str:
    """
    Fetches a popular meme related to the given topic using Giphy API.

    Args:
        topic (str): The topic to search memes for.

    Returns:
        str: URL of the fetched meme GIF.
    """
    giphy_api_key = os.getenv("GIPHY_API_KEY")
    if not giphy_api_key:
        return "Giphy API key is not configured."

    try:
        response = requests.get(
            "https://api.giphy.com/v1/gifs/search",
            params={
                "api_key": giphy_api_key,
                "q": topic,
                "limit": 1,
                "rating": "pg-13",
            }
        )
        data = response.json()
        if data["data"]:
            meme_url = data["data"][0]["images"]["downsized_medium"]["url"]
            return meme_url
        else:
            return "No memes found for the given topic."
    except Exception as e:
        print(f"Error fetching meme: {e}")
        return "Failed to fetch meme."
    

def generate_ai_image(prompt: str, size: str = "512x512") -> str:
    """
    Calls Azure OpenAI DALL·E 3 to generate an image from a prompt, using the
    asynchronous 'submit -> poll -> retrieve' workflow. Downloads the result
    to a local file and returns a /download_image/<filename> route.

    Args:
        prompt (str): Text prompt describing the desired image.
        size (str): One of "256x256", "512x512", or "1024x1024" typically.

    Returns:
        str: A local route-based URL to the downloaded image file,
             or an error message if generation failed.
    """
    # 1. Get environment config for Azure
    #    Example: AZURE_OPENAI_ENDPOINT="https://YOUR_RESOURCE_NAME.openai.azure.com"
    #             AOAI_KEY="YOUR-AZURE-OPENAI-KEY"
    #    (Make sure you have a deployment named "dall-e-3" in that resource.)
    azure_openai_endpoint = os.getenv("AOAI_ENDPOINT")  # e.g. "https://studybuddy1.openai.azure.com"
    azure_openai_api_key = os.getenv("AOAI_KEY")
    deployment_name = "dall-e-3"  # The name of your DALL·E 3 deployment
    api_version = "2024-02-01"    # Or the version you see in the portal for your DALL·E 3

    if not azure_openai_endpoint or not azure_openai_api_key:
        return "Azure OpenAI image-generation credentials are not configured properly."

    # 2. Submit the generation request
    #    POST {endpoint}/openai/deployments/{deployment-name}/images/generations:submit?api-version=...
    #    Body must wrap prompt text in an object, e.g. "prompt": {"text": "..."}
    submit_url = (
        f"{azure_openai_endpoint}openai/deployments/{deployment_name}/images/generations:submit"
        f"?api-version={api_version}"
    )

    submit_payload = {
        "prompt": {
            "text": prompt
        },
        "n": 1,
        "size": size
    }

    try:
        submit_resp = requests.post(
            submit_url,
            headers={
                "Content-Type": "application/json",
                "api-key": azure_openai_api_key
            },
            json=submit_payload,
            timeout=60
        )
        submit_resp.raise_for_status()
    except Exception as e:
        return f"Azure OpenAI image generation (submit) failed: {e}"

    # The response will contain an operationId (often in headers or JSON)
    # We can parse from JSON or from the headers “Operation-Location” or “operationId”.
    try:
        submit_data = submit_resp.json()
        operation_id = submit_data["operationId"]  # or possibly submit_resp.headers["Operation-Location"]
    except Exception as e:
        return f"Unexpected submit response: {submit_resp.text}"

    # 3. Poll the generation status until it’s done
    status_url = (
        f"{azure_openai_endpoint}openai/images/generations:status"
        f"?api-version={api_version}&operationId={operation_id}"
    )

    for _ in range(40):  # poll up to ~40x
        try:
            stat_resp = requests.get(
                status_url,
                headers={"api-key": azure_openai_api_key},
                timeout=10
            )
            stat_resp.raise_for_status()
            stat_data = stat_resp.json()
        except Exception as e:
            return f"Error checking generation status: {e}"

        status = stat_data.get("status")
        if status == "succeeded":
            break
        elif status == "failed":
            return f"Image generation failed with response: {stat_data}"
        time.sleep(2)  # Wait briefly before re-checking

    if status != "succeeded":
        return f"Image generation did not succeed in time. Last status: {stat_data}"

    # 4. Retrieve the final result (once status is "succeeded")
    result_url = (
        f"{azure_openai_endpoint}openai/images/generations:result"
        f"?api-version={api_version}&operationId={operation_id}"
    )

    try:
        result_resp = requests.get(
            result_url,
            headers={"api-key": azure_openai_api_key},
            timeout=30
        )
        result_resp.raise_for_status()
    except Exception as e:
        return f"Could not retrieve final image result: {e}"

    result_data = result_resp.json()
    # Typically shape: { "result": { "data": [ { "url": "https://..." } ] } }
    try:
        image_url = result_data["result"]["data"][0]["url"]
    except Exception as e:
        return f"Unexpected result JSON: {result_data}"

    # 5. Download the resulting image and store locally
    try:
        img_resp = requests.get(image_url, timeout=30)
        img_resp.raise_for_status()
        filename = f"{uuid.uuid4()}.png"
        output_dir = "generated_images"
        os.makedirs(output_dir, exist_ok=True)
        file_path = os.path.join(output_dir, filename)

        with open(file_path, "wb") as f:
            f.write(img_resp.content)

        # Return a local URL to your Flask route
        backend_base_url = os.getenv("BACKEND_BASE_URL", "http://localhost:8000")
        download_url = f"{backend_base_url}/ai_mentor/download_image/{filename}"

        return download_url

    except Exception as e:
        return f"Could not download or store generated image: {str(e)}"